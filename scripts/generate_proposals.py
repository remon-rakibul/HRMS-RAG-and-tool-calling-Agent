"""
Generate project proposal DOCX files for founders.
Requires: pip install python-docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT_DIR = Path(__file__).resolve().parent.parent


def set_heading_style(doc, level, font_size_pt):
    """Ensure heading styles have consistent font size."""
    style = doc.styles[f'Heading {level}']
    style.font.size = Pt(font_size_pt)
    style.font.bold = True


def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p


def build_hrms_livekit_proposal():
    doc = Document()
    doc.add_heading('Project Proposal: HRMS Agent Voice Mode Integration with LiveKit', 0)

    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'This proposal outlines the integration of the existing HRMS Agentic AI application with LiveKit Voice AI, '
        'enabling employees and admins to interact with the HRMS via voice from within the same chat interface. '
        'A "Voice Mode" button in the chat UI will start a LiveKit-powered voice session that reuses the same '
        'LangGraph RAG backend, tool-calling, and Human-in-the-Loop (HITL) flows—with full support for '
        'interruptions and natural turn-taking.'
    )

    doc.add_heading('2. In Plain English', 1)
    doc.add_paragraph(
        'Today, employees use the HRMS by typing in a chat: "Apply for leave tomorrow" or "What is my leave balance?" '
        'With this project, we add a "Voice Mode" button. When someone clicks it, they can do the same things by speaking '
        'instead of typing. The assistant hears them, looks up the same policies and data, and can even ask for approval '
        'by voice (e.g. "Say approve or reject"). If the user interrupts while the assistant is talking, the system '
        'handles it naturally—just like a real phone call. Everything stays in one place: one chat, one conversation, '
        'whether they type or talk.'
    )

    doc.add_heading('3. User Stories', 1)
    doc.add_paragraph(
        'These are examples of what users will be able to do once Voice Mode is live:'
    )
    stories = [
        'As an employee, I want to click "Voice Mode" in the chat so that I can apply for leave or check my balance by speaking, especially when I am on the go or prefer not to type.',
        'As an admin, I want to approve leave or attendance by voice so that I can handle requests quickly without switching to the keyboard.',
        'As an employee, I want to interrupt the assistant when it is talking so that I can correct something or ask a follow-up question without waiting for it to finish.',
        'As a manager, I want the voice conversation to appear in the same chat thread as text so that my team and I have one continuous record of what was requested and approved.',
    ]
    for s in stories:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('4. Example Scenario: A Day with Voice Mode', 1)
    doc.add_paragraph(
        'Neha opens the HRMS chat on her phone during her commute. She has already been chatting with the assistant about '
        'leave. She taps "Voice Mode," and the assistant says: "You\'re now in voice mode. How can I help?" Neha says: '
        '"Apply for sick leave tomorrow." The assistant repeats the details and asks her to confirm. She says "Yes." '
        'The assistant confirms the leave is submitted. Neha asks: "What about my balance?" Before the assistant finishes '
        'reading the full balance, Neha interrupts: "Just sick leave." The assistant replies with only the sick leave '
        'balance. Neha taps "End voice" and sees the whole exchange in the same chat as text. She can continue typing '
        'or start voice again later—all in one place.'
    )

    doc.add_heading('5. Why This Matters for the Business', 1)
    benefits = [
        'More employees can use the HRMS when typing is difficult (driving, factory floor, accessibility needs).',
        'Faster handling of simple requests (leave, balance, policy questions) without opening multiple screens.',
        'Same rules and approvals in voice as in chat—no separate process to maintain.',
        'One conversation history for both text and voice, so support and audits stay simple.',
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('6. Current State (Technical Context)', 1)
    doc.add_paragraph(
        'The HRMS application today provides:'
    )
    items = [
        'Text-based chat interface (React + TypeScript) with streaming responses.',
        'LangGraph workflow backend (FastAPI) with RAG (pgvector), HRMS tool-calling (leave, attendance, approvals), and HITL approval flows.',
        'Session and thread management; chat history; JWT authentication.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        'The reference LiveKit RAG PoC (read-only repo) demonstrates a voice pipeline that uses a custom LLM adapter '
        'to call a FastAPI POST /api/v1/chat endpoint with SSE streaming. It uses LiveKit Agents for STT (e.g. Deepgram), '
        'TTS (e.g. Cartesia), VAD, turn detection, and preemptive generation—with support for natural interruptions.'
    )

    doc.add_heading('7. Proposed Solution: Voice Mode in Chat', 1)
    doc.add_paragraph(
        'Integrate LiveKit Voice AI into the existing HRMS chat product so that:'
    )
    doc.add_paragraph(
        'A "Voice Mode" (or "Start voice call") button appears inside the current chat interface. When the user clicks it, '
        'the app transitions into a voice chat mode: the user joins a LiveKit room and can speak to the same HRMS agent '
        'that powers the text chat. The agent continues to use the same backend (same /api/v1/chat endpoint, session_id, '
        'thread_id, employee context) so that RAG, tool execution, and HITL approvals remain consistent across text and voice.'
    )
    doc.add_paragraph('Key capabilities to deliver:')
    voice_items = [
        'Voice Mode button in the chat UI that starts a LiveKit voice session (same session/thread/employee context as text chat where applicable).',
        'Interruption support: user can interrupt the agent mid-response; turn-taking and VAD handled by LiveKit (e.g. MultilingualModel, Silero VAD).',
        'Preemptive generation where supported to reduce latency.',
        'STT and TTS: use production-grade models (e.g. Deepgram for STT, Cartesia or similar for TTS) as in the reference PoC.',
        'Backend reuse: extend the existing BackendLLM-style adapter to pass session_id, thread_id, and employee_id to the HRMS chat API so voice and text share the same conversation and permissions.',
        'Handling of HITL in voice: when the backend streams an interrupt (e.g. approval request), the voice pipeline can either announce the request and use a voice-menu (e.g. "Say approve or reject") or surface a short in-call prompt; optional fallback to the existing text approval card if the user is still in the same tab.',
    ]
    for item in voice_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8. Technical Approach (For the Technical Team)', 1)
    doc.add_heading('8.1 Architecture', 2)
    doc.add_paragraph(
        'Reuse the mechanism from the LiveKit RAG PoC:'
    )
    doc.add_paragraph(
        'Voice pipeline: LiveKit Agents (Python) handle STT → custom Backend LLM (calls HRMS FastAPI /api/v1/chat) → TTS. '
        'The custom LLM adapter will be extended to send session_id, thread_id, and optional employee_id in the request body '
        'so that the same RAG and tool-calling workflow runs with the correct user context. The adapter will parse SSE events '
        'type: "token", type: "done", and optionally type: "interrupt" to support HITL in voice (e.g. prompt user to confirm).'
    )
    doc.add_heading('8.2 Frontend Integration', 2)
    doc.add_paragraph(
        'Add a Voice Mode entry point in the existing React chat interface (e.g. a button in the header or message area). '
        'On click: obtain LiveKit token (from a new backend endpoint or existing auth) and room name; render an in-app voice '
        'experience (e.g. embedded LiveKit room using @livekit/components-react or a modal) so the user can talk without '
        'leaving the chat. Reuse the same session/thread so that after the voice call, the conversation continues in text '
        'with full history.'
    )
    doc.add_heading('8.3 Backend Changes', 2)
    doc.add_paragraph(
        'Minimal: the existing POST /api/v1/chat already accepts session_id, thread_id, and streams SSE. Add an endpoint '
        '(e.g. POST /api/v1/voice/token) that returns a LiveKit access token for the current user/session so the frontend '
        'can join the room securely. The LiveKit agent service will run as a separate deployable (same as in the PoC) '
        'configured with CHAT_API_URL pointing to the HRMS backend.'
    )

    doc.add_heading('9. Scope and Deliverables', 1)
    deliverables = [
        'LiveKit agent service (Python) that calls HRMS /api/v1/chat with session_id, thread_id, employee context; SSE parsing for token, done, and interrupt.',
        'Voice Mode button and in-chat voice UI (join room, leave room, optional transcript) in the existing HRMS frontend.',
        'LiveKit token endpoint in the HRMS backend for secure room access.',
        'Documentation and runbooks for deploying the LiveKit agent (e.g. LiveKit Cloud or self-hosted) and configuring STT/TTS.',
    ]
    for d in deliverables:
        doc.add_paragraph(d, style='List Number')

    doc.add_heading('10. Risks and Mitigations', 1)
    doc.add_paragraph(
        'Latency: Voice is sensitive to delay. Mitigation: use preemptive generation, low-latency STT/TTS, and keep the backend '
        'streaming path unchanged. HITL in voice: approval flows may need a short voice prompt or fallback to text; we will '
        'design one clear path (e.g. "Say approve or reject") and document it.'
    )

    doc.add_heading('11. Timeline (High-Level)', 1)
    doc.add_paragraph(
        'Phase 1: Backend LLM adapter extended for HRMS (session_id, thread_id, interrupt handling); LiveKit agent deployed and tested against HRMS backend. '
        'Phase 2: Token endpoint and Voice Mode button + in-app voice UI. Phase 3: HITL in voice and polish.'
    )

    doc.add_heading('12. Conclusion', 1)
    doc.add_paragraph(
        'Integrating LiveKit with the existing HRMS agent will provide a seamless voice experience inside the same chat product, '
        'reusing the same RAG, tools, and HITL logic. The approach is based on a proven PoC and requires no change to the '
        'core LangGraph workflow—only a new entry point (Voice Mode) and an extended adapter for context and interrupts.'
    )

    out_path = OUTPUT_DIR / 'HRMS_Agent_LiveKit_Voice_Integration_Proposal.docx'
    doc.save(str(out_path))
    return out_path


def build_recom_voice_phone_proposal():
    doc = Document()
    doc.add_heading('Project Proposal: Recom Voice AI for Phone Calls with RAG', 0)

    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'This proposal outlines the design and delivery of a Voice AI system for Recom that handles inbound and/or outbound '
        'phone calls, powered by the same RAG (Retrieval Augmented Generation) and conversational AI capabilities used in '
        'the existing HRMS agent. The system will provide accurate, document-grounded answers over the phone with natural '
        'turn-taking, interruption support, and optional handoff to human agents.'
    )

    doc.add_heading('2. In Plain English', 1)
    doc.add_paragraph(
        'When a customer or employee calls Recom, instead of (or in addition to) reaching a person, they can talk to an AI '
        'that has been trained on Recom’s own documents—policies, procedures, product information. The AI answers from that '
        'knowledge, so responses stay accurate and on-brand. Callers can interrupt to ask something else or clarify, just like '
        'in a normal conversation. If the AI cannot help, the call can be transferred to a human. All of this runs over the '
        'phone line they already use, with optional recording and logging for quality and compliance.'
    )

    doc.add_heading('3. User Stories', 1)
    doc.add_paragraph(
        'Examples of what callers and the business will get:'
    )
    stories_recom = [
        'As a customer, I want to call Recom and get answers to common questions (e.g. policies, product info) by speaking, so that I can resolve issues without waiting in a queue.',
        'As a manager, I want the AI to use only company-approved documents when answering, so that callers receive consistent, correct information.',
        'As a caller, I want to interrupt the AI when it is talking so that I can ask a follow-up or correct a misunderstanding without starting over.',
        'As the business, I want calls to be logged and optionally recorded so that we can review quality and meet compliance requirements.',
        'As the business, I want the option to transfer complex or sensitive calls to a human agent so that we never leave callers without a path to a person when needed.',
    ]
    for s in stories_recom:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('4. Example: A Customer Calls', 1)
    doc.add_paragraph(
        'A customer dials Recom’s support number. The AI greets them and asks how it can help. The customer says: "What is '
        'your return policy?" The AI looks up the return policy from Recom’s ingested documents and reads back a short, clear '
        'summary. The customer interrupts: "What about digital products?" The AI immediately answers with the digital-product '
        'section. The customer says they want to speak to someone. The AI confirms and transfers the call to an agent. The '
        'call is logged with a short summary and optional recording for training and compliance.'
    )

    doc.add_heading('5. Why This Matters for Recom', 1)
    benefits_recom = [
        'Customers get instant answers to common questions, reducing wait times and call handling load.',
        'Answers are grounded in Recom’s own documents, so the message is consistent and up to date.',
        'The same system can scale to many callers without adding headcount for simple queries.',
        'Optional handoff to humans keeps complex or sensitive cases with the team while the AI handles the rest.',
        'Recording and logging support quality assurance and compliance.',
    ]
    for b in benefits_recom:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('6. Objectives', 1)
    objectives = [
        'Enable customers or employees to call in and get answers from company knowledge (policies, procedures, product info) via voice.',
        'Use RAG over ingested documents so responses are accurate and up to date.',
        'Support natural conversation: interruptions, turn-taking, and clear prompts (e.g. "Say yes to confirm").',
        'Optional: integrate with existing telephony (e.g. SIP trunk) and CRM or ticketing for handoff and logging.',
    ]
    for o in objectives:
        doc.add_paragraph(o, style='List Bullet')

    doc.add_heading('7. Proposed Architecture (For the Technical Team)', 1)
    doc.add_paragraph(
        'Leverage the same patterns as the HRMS agent and the LiveKit RAG PoC:'
    )
    doc.add_paragraph(
        'Telephony: LiveKit supports SIP trunks, so phone calls can be routed into a LiveKit room. The same LiveKit Agents '
        'voice pipeline (STT → LLM → TTS) used for in-app voice can serve phone callers. A custom Backend LLM adapter '
        'calls a dedicated RAG API (e.g. POST /api/v1/chat or a dedicated /api/v1/voice/chat) that runs the same RAG and '
        'conversation logic over Recom’s ingested documents.'
    )
    doc.add_paragraph(
        'RAG backend: Reuse or adapt the existing LangGraph RAG pipeline (FastAPI, pgvector, document ingestion) so that '
        'the voice agent has access to the same knowledge base. Responses remain grounded in company documents and can '
        'cite or summarize policies and procedures.'
    )
    doc.add_paragraph(
        'Phone-specific: Use LiveKit’s SIP plugin to connect to a SIP trunk (e.g. Twilio, Telnyx, or Recom’s current '
        'provider). Noise cancellation (e.g. BVC Telephony) can be applied for better call quality. Optional: transfer to '
        'human agent or log the call in a CRM via webhooks or API calls from the agent.'
    )

    doc.add_heading('8. Key Features', 1)
    features = [
        'Inbound (and optionally outbound) phone calls via SIP/LiveKit.',
        'RAG over Recom’s document corpus for accurate, up-to-date answers.',
        'Interruption support and natural turn-taking (VAD, turn detection).',
        'Configurable voice (TTS) and language (STT) to match Recom’s audience.',
        'Optional: handoff to human agent, call recording, and logging for compliance.',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('9. Scope and Deliverables', 1)
    deliverables_recom = [
        'RAG-backed chat API suitable for voice (streaming SSE, low latency) and document ingestion pipeline for Recom’s knowledge base.',
        'LiveKit agent service that connects to this RAG API and handles voice (STT/TTS, turn-taking, interruptions).',
        'SIP integration so that phone calls reach the LiveKit room (inbound; outbound if in scope).',
        'Operational runbooks: deployment (LiveKit Cloud or self-hosted), SIP trunk configuration, and monitoring.',
    ]
    for d in deliverables_recom:
        doc.add_paragraph(d, style='List Number')

    doc.add_heading('10. Risks and Mitigations', 1)
    doc.add_paragraph(
        'Phone quality: Network and handset variability can affect STT. Mitigation: use robust STT (e.g. Deepgram), '
        'noise cancellation, and clear prompts. Compliance: ensure call recording and data handling meet Recom’s legal '
        'and privacy requirements; design logging and retention up front.'
    )

    doc.add_heading('11. Conclusion', 1)
    doc.add_paragraph(
        'A Recom Voice AI for phone calls, built on the same RAG and LiveKit voice patterns as the HRMS agent, will '
        'provide consistent, document-grounded support over the phone. The proposal reuses proven components (LiveKit '
        'Agents, Backend LLM adapter, LangGraph RAG) and focuses on SIP integration, Recom-specific content, and '
        'operational readiness.'
    )

    out_path = OUTPUT_DIR / 'Recom_Voice_AI_Phone_Calls_RAG_Proposal.docx'
    doc.save(str(out_path))
    return out_path


if __name__ == '__main__':
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    p1 = build_hrms_livekit_proposal()
    p2 = build_recom_voice_phone_proposal()
    print(f'Generated: {p1}')
    print(f'Generated: {p2}')
